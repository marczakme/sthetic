import os
import re
import json
import time
from io import BytesIO
from dataclasses import dataclass
from typing import List, Dict, Tuple, Optional

import streamlit as st

# OpenAI Python SDK (v1+)
from openai import OpenAI
from openai import APIError, RateLimitError, APITimeoutError, APIConnectionError

# DOCX
from docx import Document
from docx.shared import Pt


# =========================
# Konfiguracja
# =========================

DEFAULT_MODEL = "gpt-4.1-mini"
MAX_RETRIES = 4

# „Agresywniej”, ale nadal: 0 nowych faktów medycznych.
SYSTEM_PROMPT = """Jesteś redaktorem medycznym oraz specjalistą SEO (Surfer SEO).
Twoim zadaniem jest zoptymalizować artykuł medyczny w Markdown pod wskazane frazy.

KRYTYCZNE ZASADY RZETELNOŚCI (NIE DO NARUSZENIA):
- Opieraj się WYŁĄCZNIE na informacjach zawartych w dostarczonym tekście.
- Nie dodawaj żadnych nowych faktów medycznych, zaleceń, statystyk, diagnoz, przeciwwskazań ani informacji klinicznych.
- Nie „dopowiadaj” wiedzy ogólnej. Jeśli czegoś nie ma w tekście, nie wolno tego wprowadzić.
- Jeśli fraza nie pasuje merytorycznie do treści lub wymagałaby dodania nowych informacji, pomiń ją.

ZASADY SEO / SURFER (AGRESYWNA, ALE BEZPIECZNA OPTYMALIZACJA):
- Priorytet: użyj jak największej liczby fraz w DOKŁADNYM brzmieniu (literalnie).
- Możesz zwiększać częstotliwość fraz przez:
  * krótkie doprecyzowania zdań w istniejących akapitach,
  * krótkie śródtytuły H3 (bez zmiany sensu),
  * sekcje podsumowań i FAQ na końcu, które powtarzają i parafrazują treść WYŁĄCZNIE z artykułu.
- Unikaj keyword stuffingu w jednym miejscu: rozkładaj frazy równomiernie.

STRUKTURA:
- Zachowaj format Markdown i istniejącą strukturę nagłówków.
- Nie przestawiaj sekcji.
- Nowe sekcje H2 możesz dodać WYŁĄCZNIE NA KOŃCU dokumentu.
- Nowe H3 w środku dokumentu są dozwolone tylko, jeśli logicznie pasują do istniejącej sekcji.

FORMAT ODPOWIEDZI:
- Zwróć WYŁĄCZNIE poprawny JSON (bez dodatkowego tekstu).
"""

USER_PROMPT_TEMPLATE = """Masz:
1) listę fraz kluczowych (każda w osobnej linii)
2) artykuł medyczny w Markdown

Cel: podnieś scoring Surfer SEO możliwie mocno, ALE bez dodawania nowych faktów medycznych.

Instrukcja obowiązkowa (dosłownie):
"Opieraj się wyłącznie na informacjach zawartych w dostarczonym tekście. Nie dodawaj nowych faktów medycznych. Jeśli fraza nie pasuje merytorycznie do treści, pomiń ją."

Zasady wykonania:
1) Najpierw wpleć frazy w istniejące akapity, tak aby były naturalne i merytoryczne.
2) Jeśli zostały frazy, które można dodać bez nowych faktów:
   - dodaj na końcu dokumentu nowe sekcje H2:
     a) "Podsumowanie"
     b) "FAQ"
   Treść w tych sekcjach ma WYŁĄCZNIE powtarzać/parafrazować informacje już obecne w artykule.
   W FAQ formułuj pytania i odpowiedzi tak, aby nie wprowadzać żadnych nowych informacji.
3) Staraj się użyć jak największej liczby fraz w dokładnym brzmieniu.
4) Kontroluj intensywność:
   - target_used_ratio: {target_used_ratio} (np. 0.85 oznacza „postaraj się użyć 85% fraz” jeśli pasują)
   - max_total_phrase_occurrences: {max_total_phrase_occurrences} (maksymalna łączna liczba wystąpień fraz – rozkładaj je po tekście)

Zwróć WYŁĄCZNIE JSON o polach:
- optimized_markdown: string
- used_phrases: array[string] (frazy faktycznie użyte, dokładnie jak wejście)
- skipped_phrases: array[string]
- brief_notes: array[string] (krótkie notatki dlaczego coś pominięto / jak rozłożono frazy; bez wiedzy medycznej)

FRAZY:
{phrases}

TEKST MARKDOWN:
{markdown}
"""


# =========================
# Helpers: frazy, walidacje, liczniki
# =========================

def normalize_phrase_lines(raw: str) -> List[str]:
    lines = [ln.strip() for ln in (raw or "").splitlines()]
    lines = [ln for ln in lines if ln]
    seen = set()
    out = []
    for ln in lines:
        key = ln.lower()
        if key not in seen:
            out.append(ln)
            seen.add(key)
    return out


def markdown_code_fences_balanced(md: str) -> bool:
    return (md.count("```") % 2) == 0


def extract_h2_headings(md: str) -> List[str]:
    headings = []
    for m in re.finditer(r"(?m)^\s*##\s+(.+?)\s*$", md):
        headings.append(m.group(1).strip())
    return headings


def validate_h2_added_only_at_end(original_md: str, optimized_md: str) -> Tuple[bool, str]:
    orig_h2 = extract_h2_headings(original_md)
    opt_h2 = extract_h2_headings(optimized_md)

    if not orig_h2:
        return True, "Brak H2 w oryginale — pomijam twardą walidację kolejności H2."

    # oryginalne H2 muszą wystąpić w tej samej kolejności
    seq_pos = 0
    for h in opt_h2:
        if seq_pos < len(orig_h2) and h == orig_h2[seq_pos]:
            seq_pos += 1
        else:
            # jeśli jeszcze nie domknęliśmy oryginalnych H2, a trafiamy na inny H2 => nowe H2 w środku
            if seq_pos < len(orig_h2):
                return False, "Wynik wstawia nowe H2 w środku dokumentu. Nowe H2 mogą być tylko na końcu."

    if seq_pos != len(orig_h2):
        return False, "Wynik narusza obecność/ciągłość oryginalnych H2 (zmienione/usunięte)."

    return True, "Walidacja H2 OK (nowe H2 — jeśli są — znajdują się na końcu)."


def count_phrase_occurrences(phrases: List[str], text: str) -> Dict[str, int]:
    """
    Liczy literalne wystąpienia frazy (case-insensitive).
    Uwaga: to celowe — Surfer często liczy literalnie.
    """
    low = text.lower()
    out = {}
    for ph in phrases:
        # prosty count; jeśli potrzebujesz boundary, można rozbudować regexem
        out[ph] = low.count(ph.lower())
    return out


# =========================
# OpenAI: call + retry + JSON parse
# =========================

@dataclass
class OptimizeResult:
    optimized_markdown: str
    used_phrases: List[str]
    skipped_phrases: List[str]
    brief_notes: List[str]
    raw_json: Dict


def get_openai_client() -> OpenAI:
    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError("Brak OPENAI_API_KEY w zmiennych środowiskowych.")
    return OpenAI(api_key=api_key)


def call_openai_with_retries(
    client: OpenAI,
    model: str,
    system_prompt: str,
    user_prompt: str,
    temperature: float = 0.3,
) -> str:
    last_err: Optional[Exception] = None

    for attempt in range(1, MAX_RETRIES + 1):
        try:
            resp = client.chat.completions.create(
                model=model,
                temperature=temperature,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
            )
            return resp.choices[0].message.content or ""
        except (RateLimitError, APITimeoutError, APIConnectionError, APIError) as e:
            last_err = e
            time.sleep(min(2 ** attempt, 16))

    raise RuntimeError(f"Nie udało się wykonać zapytania do API po {MAX_RETRIES} próbach. Ostatni błąd: {last_err}")


def parse_strict_json(text: str) -> Dict:
    text = (text or "").strip()

    try:
        return json.loads(text)
    except Exception:
        pass

    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end != -1 and end > start:
        return json.loads(text[start:end + 1])

    raise ValueError("Nie udało się sparsować JSON z odpowiedzi modelu.")


def optimize_markdown_medical(
    phrases: List[str],
    markdown: str,
    model: str,
    temperature: float,
    target_used_ratio: float,
    max_total_phrase_occurrences: int,
) -> OptimizeResult:
    client = get_openai_client()

    user_prompt = USER_PROMPT_TEMPLATE.format(
        phrases="\n".join(phrases),
        markdown=markdown,
        target_used_ratio=f"{target_used_ratio:.2f}",
        max_total_phrase_occurrences=str(max_total_phrase_occurrences),
    )

    raw = call_openai_with_retries(
        client=client,
        model=model,
        system_prompt=SYSTEM_PROMPT,
        user_prompt=user_prompt,
        temperature=temperature,
    )

    data = parse_strict_json(raw)

    optimized = str(data.get("optimized_markdown", "")).strip()
    used = data.get("used_phrases", []) or []
    skipped = data.get("skipped_phrases", []) or []
    notes = data.get("brief_notes", []) or []

    if not isinstance(used, list):
        used = []
    if not isinstance(skipped, list):
        skipped = []
    if not isinstance(notes, list):
        notes = []

    return OptimizeResult(
        optimized_markdown=optimized,
        used_phrases=[str(x) for x in used],
        skipped_phrases=[str(x) for x in skipped],
        brief_notes=[str(x) for x in notes],
        raw_json=data,
    )


# =========================
# Markdown -> DOCX (proste, ale skuteczne pod Surfer)
# =========================

def markdown_to_docx_bytes(md: str) -> bytes:
    """
    Minimalny konwerter Markdown -> DOCX:
    - #, ##, ### -> nagłówki (level 1/2/3)
    - listy "- " i "* " -> listy punktowane
    - reszta -> akapity
    Zachowuje czytelność nagłówków dla Surfera.
    """
    doc = Document()

    # domyślna czcionka (opcjonalnie)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = (md or "").splitlines()

    current_list = False
    for line in lines:
        ln = line.rstrip()

        # pusta linia -> nowy akapit (przerwa)
        if not ln.strip():
            doc.add_paragraph("")
            current_list = False
            continue

        # Nagłówki
        m1 = re.match(r"^\s*#\s+(.+)$", ln)
        m2 = re.match(r"^\s*##\s+(.+)$", ln)
        m3 = re.match(r"^\s*###\s+(.+)$", ln)

        if m1:
            doc.add_heading(m1.group(1).strip(), level=1)
            current_list = False
            continue
        if m2:
            doc.add_heading(m2.group(1).strip(), level=2)
            current_list = False
            continue
        if m3:
            doc.add_heading(m3.group(1).strip(), level=3)
            current_list = False
            continue

        # Lista punktowana
        if re.match(r"^\s*[-*]\s+.+", ln):
            text = re.sub(r"^\s*[-*]\s+", "", ln).strip()
            doc.add_paragraph(text, style="List Bullet")
            current_list = True
            continue

        # Normalny akapit (bez agresywnego formatowania inline)
        doc.add_paragraph(ln)
        current_list = False

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


# =========================
# Streamlit UI
# =========================

def main():
    st.set_page_config(page_title="Surfer SEO – Medical Optimizer (DOCX)", layout="wide")

    st.title("Surfer SEO – Medical Optimizer (Markdown in → DOCX out)")
    st.caption(
        "Wersja agresywniejsza pod Surfer: literalne frazy + kontrolowane powtórzenia + Podsumowanie/FAQ na końcu "
        "(wyłącznie na bazie treści źródłowej, bez nowych faktów medycznych)."
    )

    with st.sidebar:
        st.header("Ustawienia")
        model = st.text_input("Model OpenAI", value=DEFAULT_MODEL)
        temperature = st.slider("Temperature", 0.0, 0.9, 0.35, 0.05,
                                help="Trochę wyżej niż wcześniej: więcej redakcji/rozbudowy, ale wciąż kontrolowanie.")
        st.markdown("---")
        st.subheader("Agresywność (Surfer)")
        target_used_ratio = st.slider("Cel użycia fraz (ratio)", 0.50, 1.00, 0.85, 0.05,
                                      help="Ile fraz model ma próbować użyć (o ile pasują merytorycznie).")
        max_total_phrase_occurrences = st.number_input(
            "Maks. łączna liczba wystąpień fraz",
            min_value=10,
            max_value=500,
            value=120,
            step=10,
            help="Limit bezpieczeństwa przed totalnym keyword stuffing. Frazy będą rozkładane po tekście."
        )
        st.markdown("---")
        st.markdown("**Wymagane:** `OPENAI_API_KEY` w zmiennych środowiskowych.")

    col1, col2 = st.columns(2, gap="large")

    with col1:
        st.subheader("1) Frazy z Surfer SEO (1 fraza na linię)")
        phrases_raw = st.text_area(
            "Frazy",
            height=260,
            placeholder="np.\nobjawy alergii\nleczenie alergii\n...",
            label_visibility="collapsed",
        )
        phrases = normalize_phrase_lines(phrases_raw)
        st.write(f"Frazy po czyszczeniu: **{len(phrases)}**")

    with col2:
        st.subheader("2) Tekst wejściowy (Markdown)")
        markdown_in = st.text_area(
            "Markdown wejściowy",
            height=260,
            placeholder="Wklej artykuł w Markdown...",
            label_visibility="collapsed",
        )
        st.write(f"Długość tekstu: **{len(markdown_in)}** znaków")

    st.markdown("---")
    run = st.button("🚀 Optymalizuj (agresywnie)", type="primary", use_container_width=True,
                    disabled=(not phrases or not markdown_in))

    if run:
        if not markdown_code_fences_balanced(markdown_in):
            st.error("Wejściowy Markdown ma niezamknięty blok ``` (nieparzysta liczba). Popraw i spróbuj ponownie.")
            st.stop()

        try:
            with st.status("Optymalizuję treść pod Surfer…", expanded=True) as status:
                status.write("Łączenie z API OpenAI…")
                result = optimize_markdown_medical(
                    phrases=phrases,
                    markdown=markdown_in,
                    model=model,
                    temperature=temperature,
                    target_used_ratio=target_used_ratio,
                    max_total_phrase_occurrences=int(max_total_phrase_occurrences),
                )

                status.write("Walidacja struktury H2…")
                ok_h2, msg_h2 = validate_h2_added_only_at_end(markdown_in, result.optimized_markdown)
                if not ok_h2:
                    st.error(msg_h2)
                    st.info("Wskazówka: zmniejsz temperature lub ogranicz agresywność (ratio / max occurrences).")
                    st.stop()
                else:
                    status.write(msg_h2)

                status.write("Liczenie wystąpień fraz…")
                occ = count_phrase_occurrences(phrases, result.optimized_markdown)

                status.write("Generowanie DOCX…")
                docx_bytes = markdown_to_docx_bytes(result.optimized_markdown)

                status.update(label="Gotowe ✅", state="complete", expanded=False)

        except Exception as e:
            st.error(f"Błąd: {e}")
            st.stop()

        # Raport: wykryte użycia literalne
        used_detected = [p for p, c in occ.items() if c > 0]
        unused_detected = [p for p, c in occ.items() if c == 0]
        total_occ = sum(occ.values())

        st.subheader("3) Wynik: podgląd Markdown (roboczy)")
        st.text_area("Zoptymalizowany Markdown", value=result.optimized_markdown, height=320)

        st.subheader("4) Pobierz plik DOCX (do Surfera)")
        st.download_button(
            "💾 Pobierz optimized.docx",
            data=docx_bytes,
            file_name="optimized.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

        st.markdown("---")
        st.subheader("5) Licznik fraz (Surfer-oriented)")

        m1, m2, m3, m4 = st.columns(4)
        m1.metric("Frazy wejściowe", len(phrases))
        m2.metric("Użyte (wykryte literalnie)", len(used_detected))
        m3.metric("Nieużyte (wykryte)", len(unused_detected))
        m4.metric("Łączne wystąpienia fraz", total_occ)

        with st.expander("Wystąpienia każdej frazy (literalnie)"):
            # sort: najpierw najczęstsze
            rows = sorted(occ.items(), key=lambda x: x[1], reverse=True)
            for ph, c in rows:
                st.write(f"- **{ph}** → {c}")

        with st.expander("Deklaracja modelu: used_phrases / skipped_phrases"):
            st.markdown("**used_phrases (deklaracja modelu):**")
            st.write(result.used_phrases)
            st.markdown("**skipped_phrases (deklaracja modelu):**")
            st.write(result.skipped_phrases)

        if result.brief_notes:
            with st.expander("Notatki modelu (brief_notes)"):
                for n in result.brief_notes:
                    st.write(f"- {n}")

        with st.expander("Debug: surowy JSON z modelu"):
            st.json(result.raw_json)


if __name__ == "__main__":
    main()
